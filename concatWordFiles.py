import os
import argparse
from docx import Document
from tqdm import tqdm

def list_docx_files(folder):
    """Return sorted list of .docx file names in folder."""
    return sorted([
        f for f in os.listdir(folder)
        if f.lower().endswith('.docx') and os.path.isfile(os.path.join(folder, f))
    ])

def apply_title_style(doc, title_text):
    """Insert title with preferred Word style."""
    paragraph = doc.add_paragraph(title_text)
    for style_choice in ['Heading 1', 'Title']:
        try:
            paragraph.style = style_choice
            break
        except Exception:
            continue  # Try next style

def count_units(text, unit):
    """Count number of words or sentences in a given text."""
    if unit == 'words':
        return len(text.split())
    elif unit == 'sentences':
        return text.count('.') + text.count('!') + text.count('?')
    return 0

def merge_word_files(folder_path='.', list_only=False, progress_unit='sentences'):
    """Concatenate Word documents with progress and styled file titles."""
    docx_files = list_docx_files(folder_path)

    if list_only:
        print("\n🗂 Files to be processed:")
        for name in docx_files:
            print(f"- {name}")
        return

    merged = Document()

    for file_name in docx_files:
        file_path = os.path.join(folder_path, file_name)
        sub_doc = Document(file_path)

        apply_title_style(merged, file_name)

        # Create progress bar per file
        total_units = sum(count_units(p.text, progress_unit) for p in sub_doc.paragraphs)
        pbar = tqdm(total=total_units, desc=f"🔄 Merging '{file_name}'", unit=progress_unit)

        for para in sub_doc.paragraphs:
            merged.add_paragraph(para.text, style=para.style)
            pbar.update(count_units(para.text, progress_unit))
        pbar.close()

    output_path = os.path.join(folder_path, "merged_output.docx")
    merged.save(output_path)
    print(f"\n✅ Output saved to: {output_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Concatenate Word files in a folder.")
    parser.add_argument("folder", nargs="?", default=".", help="Target folder path (default: current)")
    parser.add_argument("--list", action="store_true", help="Only list the files to be processed")
    parser.add_argument("--unit", choices=["words", "sentences"], default="sentences", help="Progress tracking unit")
    args = parser.parse_args()

    merge_word_files(args.folder, list_only=args.list, progress_unit=args.unit)